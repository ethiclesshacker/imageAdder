#! python3

import docx
from docx.shared import Mm
import sys

pageWidth = 210
pageHeight = 297
marginSize = 12.7

doc = docx.Document()
section = doc.sections[0]
section.page_height = Mm(pageHeight)
section.page_width = Mm(pageWidth)
section.left_margin = Mm(marginSize)
section.right_margin = Mm(marginSize)
section.top_margin = Mm(marginSize)
section.bottom_margin = Mm(marginSize)

width = ((pageWidth - (marginSize * 2)) / 2) - 5
height = ((pageHeight - (marginSize * 2)) / 3) - 5


# Replace the sys.argv with a proper array when doing __name__ == __main__

run = doc.add_paragraph().add_run()

for i in range(1,(len(sys.argv)),2):
    run.add_picture(sys.argv[i],Mm(width),Mm(height))
    run.add_text("      ")    
    if i == (len(sys.argv)-1):
        break
    run.add_picture(sys.argv[i+1],Mm(width),Mm(height))
    run.add_break()
    run.add_break()

doc.save("images.docx")